"""
SAM.gov Contract Data Extractor -- Multi-Document Version
=========================================================
An enhanced version of the extractor that processes an entire solicitation
package (multiple files in a folder) rather than a single document.

Real-world solicitations on SAM.gov often include multiple attachments:
- The SAM.gov synopsis/notice
- Performance Work Statement (PWS) or Statement of Work (SOW)
- Section L -- Instructions to Offerors
- Section M -- Evaluation Criteria
- Pricing template or CLIN structure (Excel)
- Contract Data Requirements List (CDRLs)
- Attachments (org charts, security requirements, past performance forms)

Supported file types: .txt, .pdf, .docx, .xlsx, .xls

Required libraries (install once):
    pip install anthropic pdfplumber python-docx openpyxl

Built to demonstrate government-relevant AI use cases using the Anthropic API.

Author: Harry Cotton
"""

import anthropic
import sys
import os
import json
import glob


# ============================================================
# STEP 1: FILE TYPE EXTRACTORS
# ============================================================

def extract_text_from_txt(file_path: str) -> str:
    with open(file_path, "r", encoding="utf-8") as f:
        return f.read()


def extract_text_from_pdf(file_path: str) -> str:
    try:
        import pdfplumber
    except ImportError:
        print("Error: pdfplumber is required for PDF files. Run: pip install pdfplumber")
        sys.exit(1)

    text = ""
    with pdfplumber.open(file_path) as pdf:
        for i, page in enumerate(pdf.pages):
            page_text = page.extract_text()
            if page_text:
                text += f"[Page {i + 1}]\n{page_text}\n\n"
    return text


def extract_text_from_docx(file_path: str) -> str:
    try:
        from docx import Document
    except ImportError:
        print("Error: python-docx is required for Word files. Run: pip install python-docx")
        sys.exit(1)

    doc = Document(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))

    return "\n".join(paragraphs)


def extract_text_from_excel(file_path: str) -> str:
    try:
        import openpyxl
    except ImportError:
        print("Error: openpyxl is required for Excel files. Run: pip install openpyxl")
        sys.exit(1)

    wb = openpyxl.load_workbook(file_path, data_only=True)
    text = ""

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]

        # Collect non-empty rows
        rows = []
        for row in ws.iter_rows(values_only=True):
            cells = [str(cell) if cell is not None else "" for cell in row]
            if any(c.strip() for c in cells):
                rows.append(cells)

        if not rows:
            continue

        text += f"[Sheet: {sheet_name}]\n"
        text += " | ".join(rows[0]) + "\n"
        text += "-" * 60 + "\n"
        for row in rows[1:]:
            text += " | ".join(row) + "\n"
        text += "\n"

    return text


# ============================================================
# STEP 2: DOCUMENT READER
# ============================================================

SUPPORTED_EXTENSIONS = {
    ".txt":  extract_text_from_txt,
    ".pdf":  extract_text_from_pdf,
    ".docx": extract_text_from_docx,
    ".xlsx": extract_text_from_excel,
    ".xls":  extract_text_from_excel,
}

OUTPUT_SUFFIXES = (
    "_extraction.txt",
    "_extracted.txt",
    "_full_extraction.json",
    "_raw_extraction.txt",
)

FILE_TYPE_LABELS = {
    "txt": "Text", "pdf": "PDF", "docx": "Word", "xlsx": "Excel", "xls": "Excel"
}


def read_all_documents(folder_path: str) -> str:
    """
    Reads all supported files in a folder and combines them into a single
    labeled string. Each document is clearly marked with its filename and
    type so Claude can reference specific source documents in its extraction.

    Previously generated output files are excluded automatically.
    """

    all_files = []
    for ext, extractor in SUPPORTED_EXTENSIONS.items():
        for f in glob.glob(os.path.join(folder_path, f"*{ext}")):
            basename = os.path.basename(f)
            if not any(basename.endswith(suffix) for suffix in OUTPUT_SUFFIXES):
                all_files.append((f, extractor))

    all_files.sort(key=lambda x: x[0])

    if not all_files:
        print(f"Error: No supported documents found in {folder_path}")
        print(f"Supported types: {', '.join(SUPPORTED_EXTENSIONS.keys())}")
        sys.exit(1)

    print(f"Found {len(all_files)} document(s):")
    for f, _ in all_files:
        print(f"  - {os.path.basename(f)}")
    print()

    combined_text = ""
    for file_path, extractor in all_files:
        filename = os.path.basename(file_path)
        ext = os.path.splitext(filename)[1].lower().lstrip(".")
        file_type = FILE_TYPE_LABELS.get(ext, ext.upper())

        print(f"  Reading {file_type}: {filename}...")
        content = extractor(file_path)

        combined_text += f"\n{'='*60}\n"
        combined_text += f"DOCUMENT: {filename} [{file_type}]\n"
        combined_text += f"{'='*60}\n\n"
        combined_text += content
        combined_text += "\n\n"

    print()
    return combined_text


# ============================================================
# STEP 3: EXTRACT -- Send to Claude and get structured JSON
# ============================================================

def extract_contract_data(folder_path: str) -> str:
    """
    Reads all documents in a solicitation package folder and extracts
    structured fields using Claude's API.

    The prompt instructs Claude to synthesize information across multiple
    documents, noting which document each data point came from. This mirrors
    how a BD analyst would work -- cross-referencing the synopsis with the
    PWS, checking Section M against Section L, flagging discrepancies, etc.
    """

    combined_text = read_all_documents(folder_path)

    client = anthropic.Anthropic()

    system_prompt = """You are a senior government contracts Business Development
analyst with 15 years of experience in federal procurement. You specialize in
rapidly qualifying contract opportunities from SAM.gov by extracting key data
points that inform bid/no-bid decisions.

You understand FAR/DFARS procurement terminology, NAICS codes, contract types
(FFP, T&M, CPFF, CPAF, IDIQ, BPA, etc.), set-aside categories (8(a), SDVOSB,
HUBZone, WOSB, etc.), and evaluation methodologies (LPTA, best value,
tradeoff analysis).

You are being given a COMPLETE SOLICITATION PACKAGE consisting of multiple
documents. These may include the SAM.gov synopsis, a PWS/SOW, evaluation
criteria, instructions to offerors, pricing templates, and other attachments.

Guidelines:
- Synthesize information across ALL documents in the package
- If the same field appears in multiple documents with different values,
  flag the discrepancy in capital letters
- Note which source document each key data point comes from when it adds clarity
- Extract only what is explicitly stated -- do not infer or assume
- If a field is not found in ANY of the provided documents, mark it as "Not specified"
- Flag any inconsistencies or contradictions between documents
- Be precise with FAR/DFARS terminology, dates, and dollar figures"""

    user_prompt = f"""Analyze the following federal contract opportunity package from
SAM.gov. This package contains multiple documents from the same solicitation.
Extract the key fields into a structured format by synthesizing information
across all provided documents.

Return your response as a JSON object with the following structure:

{{
    "opportunity_overview": {{
        "title": "Full opportunity title",
        "solicitation_number": "Solicitation or notice number",
        "notice_type": "e.g., Solicitation, Sources Sought, RFI, Pre-Solicitation, Award Notice",
        "agency": "Contracting agency name",
        "office": "Contracting office",
        "posted_date": "Date posted",
        "response_deadline": "Response due date and time, including timezone",
        "source_documents": "List of documents included in this package"
    }},
    "contract_details": {{
        "naics_code": "NAICS code and description",
        "psc_code": "Product/Service Code if available",
        "contract_type": "e.g., FFP, T&M, CPFF, IDIQ, BPA",
        "set_aside": "e.g., Total Small Business, 8(a), SDVOSB, Unrestricted",
        "estimated_value": "Dollar value and what it covers (base, total, ceiling)",
        "period_of_performance": "Base period and option years",
        "place_of_performance": "Location(s) where work will be performed"
    }},
    "requirements_summary": {{
        "scope_overview": "1 sentence summary of what the government is buying",
        "key_requirements": ["List of the most important technical/functional requirements"],
        "clearance_requirements": "Security clearance level required, if any",
        "key_personnel": "Any specific roles or qualifications required",
        "deliverables": "Key deliverables or CDRLs if specified"
    }},
    "evaluation_criteria": {{
        "evaluation_method": "e.g., LPTA, Best Value Tradeoff, Lowest Price",
        "factors": ["List of evaluation factors in order of importance if stated"],
        "submission_requirements": "Page limits, format requirements, volumes required",
        "notes": "Any additional evaluation details"
    }},
    "competitive_intelligence": {{
        "incumbent": "Current contract holder if mentioned",
        "predecessor_contract": "Previous contract number if referenced",
        "estimated_competition": "Any indicators of competitive landscape",
        "notable_terms": "Unusual clauses, restrictions, or requirements worth flagging"
    }},
    "bd_action_items": {{
        "pursuit_recommendation": "Quick assessment: what type of firm is best positioned",
        "key_dates": ["List all critical dates and deadlines from across all documents"],
        "risks_and_flags": ["Any red flags, inconsistencies between documents, or concerns a BD team should investigate"],
        "cross_document_discrepancies": ["Any contradictions or inconsistencies found between the provided documents"]
    }},
    "proposal_outline": {{
        "sections": [
            {{
                "title": "Proposal section title (e.g., Technical Approach, Management Approach, Past Performance, Price/Cost)",
                "requirements_addressed": ["Key requirements from the solicitation that this section must address, by name or reference"],
                "guidance": "One sentence on what to emphasize in this section to score well"
            }}
        ],
        "win_themes": ["2-3 overarching themes or differentiators to thread throughout the entire proposal"]
    }}
}}

Important: Return ONLY the JSON object, no additional text before or after it.

---

SOLICITATION PACKAGE:

{combined_text}"""

    message = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=16000,
        system=system_prompt,
        messages=[
            {"role": "user", "content": user_prompt}
        ]
    )

    return message.content[0].text


# ============================================================
# STEP 4: MAIN -- Display results and save JSON
# ============================================================

def main():
    """
    Entry point: accepts a folder path containing solicitation documents,
    runs the extractor across all files, and outputs results.

    Usage:
        python extractor_multi.py "Sample Contracts/Peacecorps_package"

    The folder can contain any mix of .txt, .pdf, .docx, .xlsx, and .xls files.
    """

    if len(sys.argv) < 2:
        print("Usage: python extractor_multi.py <path_to_solicitation_folder>")
        print("Example: python extractor_multi.py \"Sample Contracts\\CMS_Package\"")
        print(f"\nSupported file types: {', '.join(SUPPORTED_EXTENSIONS.keys())}")
        sys.exit(1)

    folder_path = sys.argv[1]

    if not os.path.isdir(folder_path):
        print(f"Error: Folder not found: {folder_path}")
        sys.exit(1)

    print(f"Processing solicitation package: {folder_path}")
    print("-" * 60)

    raw_response = extract_contract_data(folder_path)

    # Parse and display the JSON response
    try:
        cleaned = raw_response.strip()
        if "```json" in cleaned:
            cleaned = cleaned.split("```json")[1]
        if "```" in cleaned:
            cleaned = cleaned.split("```")[0]
        cleaned = cleaned.strip()

        parsed = json.loads(cleaned)

        print("=" * 60)
        print("EXTRACTED CONTRACT DATA (Multi-Document)")
        print("=" * 60)

        overview = parsed.get("opportunity_overview", {})
        details = parsed.get("contract_details", {})
        requirements = parsed.get("requirements_summary", {})
        eval_criteria = parsed.get("evaluation_criteria", {})
        intel = parsed.get("competitive_intelligence", {})
        actions = parsed.get("bd_action_items", {})
        outline = parsed.get("proposal_outline", {})

        print(f"\n  OPPORTUNITY: {overview.get('title', 'N/A')}")
        print(f"  SOLICITATION #: {overview.get('solicitation_number', 'N/A')}")
        print(f"  AGENCY: {overview.get('agency', 'N/A')}")
        print(f"  NOTICE TYPE: {overview.get('notice_type', 'N/A')}")
        print(f"  RESPONSE DUE: {overview.get('response_deadline', 'N/A')}")
        print(f"  SOURCE DOCS: {overview.get('source_documents', 'N/A')}")

        print(f"\n  NAICS: {details.get('naics_code', 'N/A')}")
        print(f"  CONTRACT TYPE: {details.get('contract_type', 'N/A')}")
        print(f"  SET-ASIDE: {details.get('set_aside', 'N/A')}")
        print(f"  EST. VALUE: {details.get('estimated_value', 'N/A')}")
        print(f"  PERIOD: {details.get('period_of_performance', 'N/A')}")
        print(f"  LOCATION: {details.get('place_of_performance', 'N/A')}")

        print(f"\n  SCOPE: {requirements.get('scope_overview', 'N/A')}")
        print(f"  CLEARANCE: {requirements.get('clearance_requirements', 'N/A')}")
        print(f"  EVAL METHOD: {eval_criteria.get('evaluation_method', 'N/A')}")
        print(f"  INCUMBENT: {intel.get('incumbent', 'N/A')}")

        key_reqs = requirements.get("key_requirements", [])
        if key_reqs:
            print(f"\n  KEY REQUIREMENTS:")
            for i, req in enumerate(key_reqs, 1):
                print(f"    {i}. {req}")

        factors = eval_criteria.get("factors", [])
        if factors:
            print(f"\n  EVALUATION FACTORS:")
            for i, factor in enumerate(factors, 1):
                print(f"    {i}. {factor}")

        risks = actions.get("risks_and_flags", [])
        if risks:
            print(f"\n  RISKS / FLAGS:")
            for i, risk in enumerate(risks, 1):
                print(f"    {i}. {risk}")

        discrepancies = actions.get("cross_document_discrepancies", [])
        if discrepancies:
            print(f"\n  CROSS-DOCUMENT DISCREPANCIES:")
            for i, d in enumerate(discrepancies, 1):
                print(f"    {i}. {d}")

        win_themes = outline.get("win_themes", [])
        sections = outline.get("sections", [])
        if win_themes or sections:
            print(f"\n  PROPOSAL OUTLINE:")
            if win_themes:
                print(f"    Win Themes:")
                for theme in win_themes:
                    print(f"      - {theme}")
            for section in sections:
                print(f"\n    {section.get('title', 'Section')}:")
                reqs = section.get("requirements_addressed", [])
                if reqs:
                    print(f"      Addresses: {', '.join(reqs)}")
                guidance = section.get("guidance", "")
                if guidance:
                    print(f"      Note: {guidance}")

        print("\n" + "=" * 60)

        folder_name = os.path.basename(os.path.normpath(folder_path))
        output_path = os.path.join(folder_path, f"{folder_name}_full_extraction.json")
        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(parsed, f, indent=2)
        print(f"\nFull JSON saved to: {output_path}")

    except json.JSONDecodeError:
        print("Note: Response was not valid JSON. Showing raw output.\n")
        sys.stdout.buffer.write(raw_response.encode("utf-8", errors="replace") + b"\n")

        folder_name = os.path.basename(os.path.normpath(folder_path))
        output_path = os.path.join(folder_path, f"{folder_name}_raw_extraction.txt")
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(raw_response)
        print(f"\nRaw output saved to: {output_path}")


if __name__ == "__main__":
    main()
