"""
Compliance Matrix Generator
============================
Reads the JSON extraction output from extractor_multi.py and writes a
Section L/M compliance matrix as a formatted Word document (.docx).

Federal proposals are typically required to include a compliance matrix
that maps every Section L instruction and Section M evaluation criterion
to the proposal section addressing it. Evaluators use this table to verify
completeness; missing entries are grounds for non-compliance findings.

Usage:
    python compliance_matrix.py <extraction_json>

Example:
    python compliance_matrix.py "Sample Contracts/US_Secret Service/US_Secret Service_full_extraction.json"

Output:
    A .docx file saved alongside the extraction JSON, ready to paste into
    a proposal or submit as a standalone compliance exhibit.

Required libraries:
    pip install python-docx
"""

import sys
import os
import json
from pathlib import Path


def load_extraction(json_path: str) -> dict:
    with open(json_path, "r", encoding="utf-8") as f:
        return json.load(f)


def build_matrix_rows(extraction: dict) -> list[dict]:
    """
    Builds a flat list of matrix rows from the extraction JSON.
    Each row has: source, reference, requirement_summary, proposal_section, notes.

    Sources:
      - Section L (submission/compliance instructions)
      - Section M (evaluation factors and subfactors)
      - Key requirements from requirements_summary
    """
    rows = []
    section_l = extraction.get("section_l", {})
    section_m = extraction.get("section_m", {})
    outline = extraction.get("proposal_outline", {})
    requirements = extraction.get("requirements_summary", {})

    # Build lookups from proposal outline sections
    sections = outline.get("sections", [])
    section_titles = [s.get("title", "") for s in sections]

    req_to_section = {}
    for section in sections:
        title = section.get("title", "")
        for req in section.get("requirements_addressed", []):
            req_to_section[req.lower()] = title

    def find_section(keyword: str) -> str:
        keyword_lower = keyword.lower()

        # 1. Direct substring match against section titles (catches "Factor I", "Factor II" etc.)
        for title in section_titles:
            if keyword_lower in title.lower() or title.lower() in keyword_lower:
                return title

        # 2. Keyword-level overlap with section titles (≥2 meaningful words in common)
        kw_words = set(w for w in keyword_lower.split() if len(w) > 3)
        best_title, best_overlap = "", 0
        for title in section_titles:
            title_words = set(w for w in title.lower().split() if len(w) > 3)
            overlap = len(kw_words & title_words)
            if overlap > best_overlap:
                best_overlap, best_title = overlap, title
        if best_overlap >= 2:
            return best_title

        # 3. Substring match against requirements_addressed strings
        for req_text, section_title in req_to_section.items():
            if keyword_lower in req_text or req_text in keyword_lower:
                return section_title

        # 4. Word-overlap with requirements_addressed
        best_req_title, best_req_overlap = "", 0
        for req_text, section_title in req_to_section.items():
            req_words = set(w for w in req_text.split() if len(w) > 3)
            overlap = len(kw_words & req_words)
            if overlap > best_req_overlap:
                best_req_overlap, best_req_title = overlap, section_title
        if best_req_overlap >= 2:
            return best_req_title

        return "[TO BE CONFIRMED]"

    # --- Section L rows ---
    sub_req = section_l.get("submission_requirements", {})
    sub_req_fields = [
        ("Page Limit (Total)", sub_req.get("page_limit_total", "Not specified")),
        ("Page Limits by Volume", str(sub_req.get("page_limits_by_volume", "Not specified"))),
        ("Font Requirement", sub_req.get("font", "Not specified")),
        ("Line Spacing", sub_req.get("line_spacing", "Not specified")),
        ("File Format", sub_req.get("file_format", "Not specified")),
        ("Number of Copies", sub_req.get("copies", "Not specified")),
        ("Submission Method", sub_req.get("submission_method", "Not specified")),
    ]
    for label, value in sub_req_fields:
        if value and value not in ("Not specified", "None", "{}"):
            rows.append({
                "source": "Section L",
                "reference": "Submission Requirements",
                "requirement_summary": f"{label}: {value}",
                "proposal_section": "Cover / Transmittal Letter",
                "notes": "",
            })

    for volume in section_l.get("volume_structure", []):
        rows.append({
            "source": "Section L",
            "reference": "Volume Structure",
            "requirement_summary": f"Required volume: {volume}",
            "proposal_section": volume,
            "notes": "",
        })

    signing = section_l.get("signing_requirements", "")
    if signing and signing != "Not specified":
        rows.append({
            "source": "Section L",
            "reference": "Signing Requirements",
            "requirement_summary": signing,
            "proposal_section": "Cover / Transmittal Letter",
            "notes": "",
        })

    other = section_l.get("other_instructions", "")
    if other and other != "Not specified":
        rows.append({
            "source": "Section L",
            "reference": "Other Instructions",
            "requirement_summary": other,
            "proposal_section": "[TO BE CONFIRMED]",
            "notes": "",
        })

    # --- Section M rows ---
    for factor in section_m.get("factors", []):
        if not isinstance(factor, dict):
            continue
        factor_name = factor.get("name", "")
        weight = factor.get("weight", "Not specified")
        subfactors = factor.get("subfactors", [])

        matched_section = find_section(factor_name)

        if subfactors:
            for sub in subfactors:
                rows.append({
                    "source": "Section M",
                    "reference": f"{factor_name} (subfactor)",
                    "requirement_summary": sub,
                    "proposal_section": find_section(sub) or matched_section,
                    "notes": f"Factor weight: {weight}",
                })
        else:
            rows.append({
                "source": "Section M",
                "reference": "Evaluation Factor",
                "requirement_summary": f"{factor_name}",
                "proposal_section": matched_section,
                "notes": f"Weight: {weight}",
            })

    # --- Key requirements ---
    for req in requirements.get("key_requirements", []):
        rows.append({
            "source": "PWS/SOW",
            "reference": "Key Requirement",
            "requirement_summary": req,
            "proposal_section": find_section(req),
            "notes": "",
        })

    return rows


def write_docx(rows: list[dict], extraction: dict, output_path: str) -> None:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        print("Error: python-docx is required. Run: pip install python-docx")
        sys.exit(1)

    overview = extraction.get("opportunity_overview", {})
    title = overview.get("title", "Unknown Opportunity")
    sol_num = overview.get("solicitation_number", "N/A")
    agency = overview.get("agency", "N/A")

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Title block
    heading = doc.add_heading("Section L/M Compliance Matrix", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Opportunity: {title}\n").bold = True
    meta.add_run(f"Solicitation #: {sol_num}  |  Agency: {agency}")

    doc.add_paragraph()

    # Table
    col_headers = ["Source", "Reference", "Requirement / Instruction", "Proposal Section", "Notes"]
    col_widths = [Inches(0.8), Inches(1.3), Inches(2.6), Inches(1.8), Inches(1.0)]

    table = doc.add_table(rows=1, cols=len(col_headers))
    table.style = "Table Grid"

    # Header row
    hdr_cells = table.rows[0].cells
    for i, (header, width) in enumerate(zip(col_headers, col_widths)):
        hdr_cells[i].width = width
        hdr_cells[i].text = header
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(9)
        # Dark blue background
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "1F3864")
        tcPr.append(shd)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Data rows with alternating shading
    fill_colors = ["FFFFFF", "EEF2F7"]
    source_fill = {"Section L": "E8F0FE", "Section M": "FFF3E0", "PWS/SOW": "F1F8E9"}

    for idx, row_data in enumerate(rows):
        row_cells = table.add_row().cells
        values = [
            row_data["source"],
            row_data["reference"],
            row_data["requirement_summary"],
            row_data["proposal_section"],
            row_data["notes"],
        ]
        fill = source_fill.get(row_data["source"], fill_colors[idx % 2])

        for i, (cell, val) in enumerate(zip(row_cells, values)):
            cell.width = col_widths[i]
            cell.text = val
            run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(val)
            run.font.size = Pt(8.5)
            if row_data["proposal_section"] == "[TO BE CONFIRMED]" and i == 3:
                run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
                run.bold = True
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), fill)
            tcPr.append(shd)

    doc.save(output_path)


def main():
    if len(sys.argv) < 2:
        print("Usage: python compliance_matrix.py <extraction_json>")
        print()
        print("Example:")
        print('  python compliance_matrix.py "Sample Contracts/US_Secret Service/US_Secret Service_full_extraction.json"')
        sys.exit(1)

    json_path = sys.argv[1]

    if not os.path.exists(json_path):
        print(f"Error: Extraction JSON not found: {json_path}")
        sys.exit(1)

    print(f"Loading extraction: {json_path}")
    extraction = load_extraction(json_path)

    overview = extraction.get("opportunity_overview", {})
    print(f"Opportunity: {overview.get('title', 'Unknown')}")
    print(f"Agency:      {overview.get('agency', 'N/A')}")
    print()

    print("Building compliance matrix rows...")
    rows = build_matrix_rows(extraction)
    print(f"  {len(rows)} rows generated")
    print()

    stem = Path(json_path).stem.replace("_full_extraction", "")
    output_path = os.path.join(os.path.dirname(os.path.abspath(json_path)), f"{stem}_compliance_matrix.docx")

    print("Writing Word document...")
    write_docx(rows, extraction, output_path)

    print(f"Compliance matrix saved to: {output_path}")
    print()
    to_confirm = sum(1 for r in rows if r["proposal_section"] == "[TO BE CONFIRMED]")
    if to_confirm:
        print(f"  {to_confirm} row(s) marked [TO BE CONFIRMED] — review proposal section mapping.")


if __name__ == "__main__":
    main()
